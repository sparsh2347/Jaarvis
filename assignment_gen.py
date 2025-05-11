def generate_solution(file_path):
    #time used for adding delays when required
    import time
    #read and extract text from pdf files
    import PyPDF2
    #to create and save pdfs
    from fpdf import FPDF
    #file path opertaions
    import os
    #read and write word files
    import docx
    from docx import Document

    #selenium imports from web automation
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options

    # --- Step 1: Kill existing Chrome sessions ---
    try:
        os.system("taskkill /f /im chrome.exe")
        os.system("taskkill /f /im chromedriver.exe")
    except Exception as e:
        print("⚠️ Failed to kill existing Chrome processes:", e)

    # --- Step 2: Configure Chrome with user profile ---
    try:
        chrome_options = Options()
        chrome_options.add_argument("--user-data-dir=C:/selenium_profiles/myprofile")
        chrome_options.add_argument("profile-directory=Profile 1")  # Exactly as
        # Use your chromedriver path
        #chromedriver acts as a bridge between selenium code and actual browser
        service = Service("C:/Users/spars/OneDrive/Desktop/Google Downloads/chromedriver-win64/chromedriver-win64/chromedriver.exe")
        driver = webdriver.Chrome(service=service, options=chrome_options)
        driver.maximize_window()
    except Exception as e:
        print("❌ Error launching Chrome browser:", e)
        exit()

    # --- Step 3: Navigate to ChatGPT ---
    try:
        driver.get("https://chatgpt.com/")
        time.sleep(1)
    except Exception as e:
        print("❌ Failed to open ChatGPT website:", e)
        driver.quit()
        exit()

    # --- Step 4: Read and extract text from file ---
    def extract_text(path):
        try:
            #checking the file extension
            ext = os.path.splitext(path)[-1].lower()
            if ext == ".pdf":
                #uses PyPDF2.pdfreader to read data from pdf files
                with open(path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                #uses docx.Document to read word documents
                doc = docx.Document(path)
                return "\n".join(para.text for para in doc.paragraphs)
            elif ext == ".txt":
                #reads the entire file as plain text
                with open(path, "r", encoding="utf-8") as f:
                    return f.read()
            else:
                raise ValueError("Unsupported file format.")
        except Exception as e:
            print("❌ Error extracting text from file:", e)
            

    # --- Step 5: Save text to PDF ---
    def save_to_pdf(text,folder_path, filename):
        #initialies a new FPDF object and adds a page to it
        pdf = FPDF()
        pdf.add_page()

        #sspecifies path where the pdf would be saved
        file_sol_path = os.path.join(folder_path, filename)

        #adds a specific font to set the font of the pdf
        font_path = "C:/Users/spars/OneDrive/Desktop/Python Projects/Jaarvis/DejaVuSans.ttf"
        if not os.path.exists(font_path):
            print(f"Font file '{font_path}' not found. Please ensure it is in the same directory as the script.")
            return
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)

        #splits the text into lines and writes each line to the pdf
        for line in text.split('\n'):
            pdf.multi_cell(0, 10, line)

        try:
            #save documents in the specified path and returns the path
            pdf.output(file_sol_path)
            print(f"✅ PDF saved as {file_sol_path}")
            return file_sol_path
        except Exception as e:
            print(f"❌ Failed to save PDF: {e}")

    def save_to_docx(content,folder_path,filename):
        try:
            #initializes a Document() object from the docx module
            doc = Document()
            #adds a heading at the top
            doc.add_heading('Assignment Solution', 0)

            #formatting the content
            # Split the response by paragraphs (2 newlines)
            for para in content.strip().split("\n\n"):
                doc.add_paragraph(para.strip())
            #saves the doc to the specified folder and returns teh path
            file_doc_path = os.path.join(folder_path, filename)
            doc.save(file_doc_path)
            print(f"🎉 Solution saved to {filename}")
            return file_doc_path
        except Exception as e:
            print("❌ Failed to save DOCX:", e)


    # --- MAIN WORKFLOW ---
    try:
        # file_path = input("📂 Enter path to assignment file: ").strip('"').replace("\\", "/")
        text = extract_text(file_path)
        #Custom prompt by the user for accuarte responses
        user_prompt = input("Enter a custom prompt for GPT (or leave blank for default): ").strip()
        if user_prompt:
            final_prompt = f"{user_prompt}\n\n{text}"
        else:
            #default user prompt
            user_prompt= "Generate a detailed and accurate answer for this assignment"
            final_prompt = f"{user_prompt}\n\n{text}"
    except Exception as e:
        print("❌ Error reading input file or copying text:", e)
        driver.quit()
        exit()

    # --- Step 6: Paste into ChatGPT and submit ---
    try:
        wait=WebDriverWait(driver,20)
        # Wait for the ProseMirror input div
        editable_div = wait.until(EC.presence_of_element_located((By.XPATH, '//div[@id="prompt-textarea" and @contenteditable="true"]')))
        # Prepare the prompt text
        prompt_text = final_prompt.replace("\n", "\\n").replace('"', '\\"')

        # Insert the text directly into the ProseMirror editor using JavaScript
        driver.execute_script(f"""
            let el = arguments[0];
            el.innerHTML = '<p>{prompt_text}</p>';
            let evt = new InputEvent('input', {{ bubbles: true }});
            el.dispatchEvent(evt);
        """, editable_div)
        #By dispatching the InputEvent, the content inside the editable_div (where the assignment prompt is pasted) is marked as changed. This triggers any associated event listeners on the page, ensuring the change is recognized by the web page, and the content is updated accordingly.

        time.sleep(1)
        #presses enter after pasting the prompt into GPT input field
        driver.execute_script("""
        const event = new KeyboardEvent("keydown", {
            bubbles: true,
            cancelable: true,
            key: "Enter",
            code: "Enter",
            which: 13,
            keyCode: 13
        });
        document.querySelector('[contenteditable="true"]').dispatchEvent(event);
        """)

    except Exception as e:
        print("❌ Error pasting and submitting to ChatGPT:", e)
        driver.quit()
        exit()

    # --- Step 7: Wait for and fetch GPT response ---
    try:
        print("⌛ Waiting for GPT to finish response...")
        # Wait until at least one markdown element is present
        wait.until(EC.presence_of_element_located((By.CLASS_NAME, "markdown")))

        # Then wait until the last markdown stops changing
        prev = ""
        stable_count = 0

        while stable_count < 5:  # wait until response is stable for ~5 seconds
            time.sleep(1)
            #current repsonse
            responses = driver.find_elements(By.CLASS_NAME, "markdown")
            #last repsonse
            last = responses[-1].text.strip()
            #if last response and current response is same them stable count is increase by 1 else stable count becomes 0
            if last == prev:
                stable_count += 1
            else:
                stable_count = 0
                prev = last

        #final response after stability is confirmed
        final_response = responses[-1].text.strip()
        print("\n✅ GPT Response:\n", final_response)

    except Exception as e:
        print("❌ Failed to get GPT response:", e)
        print("⚠️ Page source for debugging:")
        driver.quit()
        exit()

    # --- Step 8: Save response to PDF ---
    try:
        folder_path="C:/Users/spars/OneDrive/Desktop/Assignements Sem 4"
        file_name=input("Enter the file name you want to save the solution as: ")
        #appends a default .docx extension to the file name
        if not file_name.endswith(".docx"):
            file_name += ".docx"
        solution_path=save_to_docx(final_response,folder_path, file_name)
        print("🎉 Solution saved to solution.pdf")
    except Exception as e:
        print("❌ Could not save GPT response to PDF:", e)
        
    # Optional: Close the browser
    driver.quit()
    
    return solution_path
